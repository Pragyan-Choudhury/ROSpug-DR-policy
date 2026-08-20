#!/usr/bin/env python3
"""
generate_architecture_doc.py
Generates an architecture Word document for:
  "Sim-to-Real Transfer with Onboard Domain Randomization — ROSPug Project"
"""

import io
import os
import matplotlib
matplotlib.use('Agg')
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
import matplotlib.patheffects as pe
import numpy as np

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── colour palette ────────────────────────────────────────────────────────────
C_INFRA   = "#2C3E50"   # dark navy   – infrastructure / docker layer
C_SIM     = "#1565C0"   # deep blue   – Gazebo simulation
C_ROS     = "#00897B"   # teal        – ROS middleware
C_ENV     = "#6A1B9A"   # purple      – RL environment
C_TRAIN   = "#E65100"   # deep orange – PPO training
C_DR      = "#C62828"   # deep red    – domain randomisation
C_EVAL    = "#2E7D32"   # dark green  – evaluation
C_REAL    = "#4527A0"   # indigo      – real hardware
C_ARROW   = "#455A64"   # slate       – arrows
C_BG      = "#FAFAFA"   # near-white  – background


# ─────────────────────────────────────────────────────────────────────────────
# 1.  ARCHITECTURE DIAGRAM
# ─────────────────────────────────────────────────────────────────────────────

def draw_box(ax, xy, w, h, label, sublabel, color, fontsize=8.5, alpha=0.92):
    """Draw a rounded-rectangle component box."""
    x, y = xy
    box = FancyBboxPatch((x, y), w, h,
                         boxstyle="round,pad=0.02",
                         linewidth=1.4,
                         edgecolor=color,
                         facecolor=color + "22",
                         zorder=3, alpha=alpha)
    ax.add_patch(box)
    # header bar
    header = FancyBboxPatch((x, y + h - 0.30), w, 0.30,
                             boxstyle="round,pad=0.01",
                             linewidth=0,
                             facecolor=color,
                             zorder=4, alpha=0.85)
    ax.add_patch(header)
    ax.text(x + w/2, y + h - 0.15, label,
            ha='center', va='center', fontsize=fontsize,
            fontweight='bold', color='white', zorder=5)
    if sublabel:
        ax.text(x + w/2, y + h/2 - 0.05, sublabel,
                ha='center', va='center', fontsize=6.8,
                color=color, zorder=5,
                linespacing=1.4)


def arrow(ax, x0, y0, x1, y1, label='', color=C_ARROW, style='->'):
    """Draw an annotated arrow."""
    ax.annotate('', xy=(x1, y1), xytext=(x0, y0),
                arrowprops=dict(arrowstyle=style,
                                color=color,
                                lw=1.4,
                                connectionstyle='arc3,rad=0.0'),
                zorder=6)
    if label:
        mx, my = (x0+x1)/2, (y0+y1)/2
        ax.text(mx, my, label, ha='center', va='bottom',
                fontsize=6.2, color=color,
                bbox=dict(boxstyle='round,pad=0.15', fc='white', ec='none', alpha=0.8),
                zorder=7)


def make_architecture_diagram() -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(16, 11))
    fig.patch.set_facecolor(C_BG)
    ax.set_facecolor(C_BG)
    ax.set_xlim(0, 16)
    ax.set_ylim(0, 11)
    ax.axis('off')

    # ── title ──────────────────────────────────────────────────────────────
    ax.text(8, 10.6, "Sim-to-Real Transfer with Domain Randomization — ROSPug Architecture",
            ha='center', va='center', fontsize=12, fontweight='bold', color=C_INFRA)

    # ── layer bands ────────────────────────────────────────────────────────
    bands = [
        (9.7, 10.45, "#2C3E50", "LAYER 1 · Development Infrastructure (Docker)"),
        (7.3,  9.6,  "#1565C0", "LAYER 2 · Simulation (Gazebo 9 / ROS Melodic)"),
        (5.2,  7.2,  "#6A1B9A", "LAYER 3 · RL Environment  (RosPugEnv)"),
        (2.9,  5.1,  "#E65100", "LAYER 4 · PPO Training  (Stable-Baselines3)"),
        (0.7,  2.8,  "#2E7D32", "LAYER 5 · Evaluation & Real-Hardware"),
    ]
    for y0, y1, c, lbl in bands:
        rect = mpatches.FancyBboxPatch((0.25, y0), 15.5, y1-y0,
                                        boxstyle="round,pad=0.05",
                                        linewidth=1.2, linestyle='--',
                                        edgecolor=c, facecolor=c+"0D", zorder=1)
        ax.add_patch(rect)
        ax.text(0.45, (y0+y1)/2, lbl,
                va='center', fontsize=6.5, color=c,
                fontweight='bold', rotation=90, zorder=2)

    # ── LAYER 1 — Docker / Infrastructure ─────────────────────────────────
    draw_box(ax, (1.1, 9.75), 4.0, 0.60,
             "Docker Container  (rospug_dev)",
             "osrf/ros:melodic-desktop-full · host network · X11",
             C_INFRA, fontsize=7.5)

    draw_box(ax, (5.5, 9.75), 3.2, 0.60,
             "Python 3.8  (deadsnakes PPA)",
             "stable-baselines3 · gymnasium\ntensorboard · rospkg",
             C_INFRA, fontsize=7.5)

    draw_box(ax, (9.1, 9.75), 2.8, 0.60,
             "Volumes / Logging",
             "rospug_research/:rw  →  logs/  checkpoints/",
             C_INFRA, fontsize=7.5)

    draw_box(ax, (12.3, 9.75), 3.2, 0.60,
             "TensorBoard  (port 6006)",
             "ep_rew_mean · ep_len_mean\napprox_kl  ·  value_loss",
             C_INFRA, fontsize=7.5)

    # ── LAYER 2 — Simulation ───────────────────────────────────────────────
    # ROSPug URDF
    draw_box(ax, (1.1, 8.60), 3.0, 0.90,
             "ROSPug URDF / XACRO",
             "12-DOF quadruped\n4 legs × (hip · thigh · calf)\nbase_link + mesh files",
             C_SIM)

    # Gazebo 9
    draw_box(ax, (4.5, 7.80), 4.2, 1.70,
             "Gazebo 9  (Physics Engine)",
             "ODE  ·  rigid-body simulation\nground friction   body mass\n"
             "servo latency  ←  DR params\n/gazebo/model_states (100 Hz)\n/gazebo/reset_world  (service)",
             C_SIM, fontsize=7.5)

    # ROS Controllers
    draw_box(ax, (9.1, 8.20), 3.5, 1.30,
             "ros_control  (12 PID controllers)",
             "/pug/{joint}_position_controller\n          /command  [Float64]\n"
             "→ effort controllers @ 50 Hz",
             C_ROS)

    # Joint States
    draw_box(ax, (12.9, 8.20), 2.6, 0.80,
             "/pug/joint_states",
             "JointState @ 100 Hz\n12 pos + 12 vel",
             C_ROS, fontsize=7.5)

    # Gait baseline
    draw_box(ax, (1.1, 7.45), 3.0, 0.85,
             "sim_gait_controller_v3.py",
             "Trot gait  ·  yaw  ·  strafe\n/cmd_vel → joint commands",
             C_SIM, fontsize=7.5)

    # ── LAYER 3 — RL Environment ───────────────────────────────────────────
    # Observation builder
    draw_box(ax, (1.1, 5.50), 3.3, 1.50,
             "Observation Builder",
             "obs[0:12]   joint positions  (rad)\n"
             "obs[12:24]  joint velocities (rad/s)\n"
             "obs[24]     body roll  (rad)\n"
             "obs[25]     body pitch  (rad)\n"
             "→  Box(−∞, ∞,  shape=(26,))",
             C_ENV)

    # RosPugEnv core
    draw_box(ax, (4.8, 5.30), 4.5, 1.70,
             "RosPugEnv   (gymnasium.Env)",
             "reset()  →  /gazebo/reset_world\n"
             "          →  stand pose  →  settle 0.5 s\n"
             "step(action)  →  pub 12 joint targets\n"
             "              →  sleep DT (0.02 s)\n"
             "              →  read obs  ·  reward  ·  done\n"
             "Ctrl rate: 50 Hz    Max steps: 500  (10 s)",
             C_ENV, fontsize=7.5)

    # Reward function
    draw_box(ax, (9.6, 5.50), 3.5, 1.50,
             "Reward Function",
             "r = vx × 3.0\n"
             "  − 0.001 × Σ(joint_vel²)\n"
             "  + 0.05  [alive bonus]\n"
             "  − 10.0  [if fallen]\n"
             "Termination:  |roll| > 0.7 rad\n"
             "              |pitch| > 0.7 rad",
             C_ENV)

    # Action space
    draw_box(ax, (13.4, 5.50), 2.1, 1.50,
             "Action Space",
             "Box(−0.5, 0.5\n shape=(12,))\nresidual offsets\nfrom stand pose\n(all zeros)",
             C_ENV, fontsize=7.5)

    # ── LAYER 4 — Training ─────────────────────────────────────────────────
    # Policy A
    draw_box(ax, (1.1, 3.10), 4.0, 1.80,
             "Policy A — Fixed Physics",
             "PPO  (MlpPolicy  [256,256])\n"
             "n_steps=2048   batch=256\n"
             "lr=3e-4   γ=0.99   λ=0.95\n"
             "clip=0.2   ent_coef=0.01\n"
             "Target: 500k–1M steps\n"
             "Checkpoint every 10k steps",
             C_TRAIN)

    # Domain Randomization
    draw_box(ax, (5.5, 3.25), 5.0, 1.65,
             "Domain Randomization Module  (Policy B)",
             "Per-episode parameter resampling:\n"
             "  Friction μ        →  U[0.4, 1.2]  (nominal 0.8)\n"
             "  Body mass         →  ±15 % of nominal\n"
             "  Servo latency     →  U[0, 50] ms\n"
             "  Service: /gazebo/set_physics_properties\n"
             "  SDF <mu1>/<mu2> tags per episode",
             C_DR, fontsize=7.8)

    # Policy B
    draw_box(ax, (10.8, 3.10), 4.7, 1.80,
             "Policy B — Domain Randomized",
             "Identical PPO hyperparameters to A\n"
             "Trains on randomised environment\n"
             "Expected: slower convergence\n"
             "Expected: better generalisation\n"
             "Separate log dir: logs/ppo_rospug_dr\n"
             "Checkpoint every 10k steps",
             C_DR)

    # ── LAYER 5 — Evaluation & Real Hardware ──────────────────────────────
    # Evaluation
    draw_box(ax, (1.1, 0.80), 6.5, 1.80,
             "Head-to-Head Evaluation  (Day 10 — Core Result)",
             "evaluate_policy.py  ·  20 episodes × 5 held-out conditions\n"
             "Held-out conditions  (neither A nor B trained on exactly):\n"
             "  μ ∈ {0.3, 0.6, 0.9, 1.3}  ·  mass ±20 %  ·  latency 60 ms\n"
             "Metrics:  fall rate  ·  distance walked  ·  episode length\n"
             "Output:   mean±std table  +  matplotlib bar chart\n"
             "Hypothesis: B more robust on held-out; A better on training cond.",
             C_EVAL, fontsize=7.5)

    # Real hardware
    draw_box(ax, (8.0, 0.80), 7.5, 1.80,
             "Stretch Goal — Real ROSPug Deployment  (Day 11–12)",
             "ROS1 inference node:  load SB3 policy → publish joint targets\n"
             "Replace simulated env action publisher\n"
             "Same /pug/{joint}_position_controller/command topics\n"
             "Safety protocol:  2–3 s burst tests first\n"
             "                  support harness  ·  torque limits\n"
             "Compare: Policy B  vs  sim_gait_controller_v3  on unseen surface",
             C_REAL, fontsize=7.5)

    # ── ARROWS ─────────────────────────────────────────────────────────────
    # URDF → Gazebo
    arrow(ax, 2.6, 8.60, 4.8, 9.00, "spawn_model", C_SIM)
    # Gazebo → ROS Controllers
    arrow(ax, 8.7, 8.85, 9.1, 8.85, "joint cmds", C_ROS)
    # ROS Controllers → joint_states
    arrow(ax, 12.6, 8.85, 12.9, 8.85, "", C_ROS)
    # joint_states → Observation Builder
    arrow(ax, 13.2, 8.20, 13.2, 7.70, "")
    arrow(ax, 13.2, 7.70, 2.75, 6.70, "/pug/joint_states", C_ROS)
    # model_states → Observation Builder
    arrow(ax, 6.9, 7.80, 2.75, 6.50, "/gazebo/model_states", C_SIM)
    # Observation Builder → RosPugEnv
    arrow(ax, 4.4, 6.25, 4.8, 6.25, "obs (26,)", C_ENV)
    # Reward → RosPugEnv
    arrow(ax, 9.6, 6.25, 9.3, 6.25, "reward/done", C_ENV)
    # Action → RosPugEnv
    arrow(ax, 13.4, 6.25, 13.2, 6.25, "action (12,)", C_ENV)
    # RosPugEnv → ros_control (joint commands)
    arrow(ax, 7.05, 7.00, 9.1, 8.20, "12 × Float64", C_ROS)
    # RosPugEnv → Gazebo reset
    arrow(ax, 5.50, 7.00, 5.20, 7.80, "/reset_world", C_SIM)
    # Policy A → RosPugEnv
    arrow(ax, 3.1, 3.10, 5.8, 5.30, "rollout\ncollect", C_TRAIN)
    # Policy B → DR → RosPugEnv
    arrow(ax, 8.0, 4.90, 7.05, 5.30, "DR params", C_DR)
    arrow(ax, 13.1, 3.10, 8.0, 5.30, "rollout\ncollect", C_DR)
    # Policies → Evaluation
    arrow(ax, 3.1, 3.10, 3.5, 2.60, "checkpoint\n.zip", C_EVAL)
    arrow(ax, 13.1, 3.10, 8.5, 2.60, "checkpoint\n.zip", C_EVAL)
    # Policy B → Real Hardware
    arrow(ax, 13.1, 3.10, 12.5, 2.60, "export\npolicy", C_REAL)
    # DR → Policy B (feedback loop label)
    arrow(ax, 8.0, 3.25, 8.0, 5.30, "reset()\nrandomize", C_DR)

    # ── legend ─────────────────────────────────────────────────────────────
    legend_items = [
        (C_INFRA,  "Infrastructure / Docker"),
        (C_SIM,    "Simulation (Gazebo)"),
        (C_ROS,    "ROS Middleware"),
        (C_ENV,    "RL Environment"),
        (C_TRAIN,  "PPO Training (Policy A)"),
        (C_DR,     "Domain Randomization (Policy B)"),
        (C_EVAL,   "Evaluation"),
        (C_REAL,   "Real Hardware"),
    ]
    lx, ly = 0.35, 0.68
    for i, (c, lbl) in enumerate(legend_items):
        col = i % 4
        row = i // 4
        px = lx + col * 3.8
        py = ly - row * 0.25
        ax.add_patch(mpatches.Rectangle((px, py-0.10), 0.25, 0.18,
                                         fc=c, ec='none', zorder=8))
        ax.text(px+0.32, py, lbl, va='center', fontsize=6.5, color='#333333', zorder=8)

    plt.tight_layout(pad=0.1)
    buf = io.BytesIO()
    plt.savefig(buf, format='png', dpi=180, bbox_inches='tight',
                facecolor=C_BG)
    plt.close(fig)
    buf.seek(0)
    return buf


# ─────────────────────────────────────────────────────────────────────────────
# 2.  WORD DOCUMENT HELPERS
# ─────────────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)


def add_heading(doc, text, level=1, color_hex="1F3864"):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(
            int(color_hex[0:2], 16),
            int(color_hex[2:4], 16),
            int(color_hex[4:6], 16))
    return p


def add_para(doc, text, bold=False, italic=False, size=10, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1F3864")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = "EBF5FB" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(9)
            set_cell_bg(cell, bg)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    return table


# ─────────────────────────────────────────────────────────────────────────────
# 3.  BUILD THE DOCUMENT
# ─────────────────────────────────────────────────────────────────────────────

def build_document(output_path: str):
    doc = Document()

    # page margins
    for section in doc.sections:
        section.top_margin    = Inches(0.9)
        section.bottom_margin = Inches(0.9)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)

    # ── Cover / Title ───────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("Sim-to-Real Transfer with Onboard Domain Randomization")
    tr.bold = True
    tr.font.size = Pt(20)
    tr.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("System Architecture Document")
    sr.font.size = Pt(13)
    sr.font.color.rgb = RGBColor(0x45, 0x5A, 0x64)

    proj = doc.add_paragraph()
    proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = proj.add_run("ROSPug Quadruped Robot   ·   ROS 1 Melodic / Gazebo 9 / Stable-Baselines3 PPO")
    pr.font.size = Pt(10)
    pr.italic = True
    pr.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr2 = date_p.add_run("Version 1.0   ·   August 2026")
    dr2.font.size = Pt(10)
    dr2.font.color.rgb = RGBColor(0x78, 0x90, 0x9C)

    doc.add_paragraph()

    # ── 1. Executive Summary ────────────────────────────────────────────────
    add_heading(doc, "1.  Executive Summary", 1)
    add_para(doc,
        "This document describes the complete software architecture for training and evaluating "
        "reinforcement-learning (RL) locomotion policies on the ROSPug 12-DOF quadruped robot. "
        "The project follows the Sim-to-Real paradigm: policies are trained entirely inside a "
        "Gazebo physics simulator and then deployed (stretch goal) to the real robot without "
        "any hardware-in-the-loop retraining.")
    add_para(doc,
        "The central research question is: does training with Domain Randomization (DR) — "
        "varying ground friction, body mass, and servo latency every episode — produce a policy "
        "that generalises better to held-out physical conditions than a baseline policy trained "
        "on fixed physics? This mirrors the approach taken in OpenAI Dactyl and ANYmal "
        "sim-to-real transfer papers, adapted to the ROSPug platform and a 2-week timeline.")

    # ── 2. Project Scope ────────────────────────────────────────────────────
    add_heading(doc, "2.  Project Scope & Deliverables", 1)
    add_table(doc,
        ["Deliverable", "Status", "Description"],
        [
            ["Policy A — Fixed Physics",    "Step 4 (planned)",  "PPO policy trained on constant Gazebo physics"],
            ["Policy B — Domain Randomized","Step 5 (planned)",  "PPO policy trained with DR (friction/mass/latency)"],
            ["Head-to-Head Evaluation",     "Step 6 (planned)",  "Fall-rate & distance table across 5 held-out conditions"],
            ["Real Hardware Deployment",    "Stretch goal",      "Policy B on real ROSPug vs fixed-gait baseline"],
            ["Environment (RosPugEnv)",     "Step 3 — DONE ✓",  "gymnasium.Env wrapper around Gazebo/ROS1"],
            ["Docker Dev Environment",      "Step 1 — DONE ✓",  "Reproducible ROS Melodic + Gazebo 9 + Python 3.8 image"],
            ["Gait Baseline Controller",    "Step 2 — DONE ✓",  "sim_gait_controller_v3.py — trot/yaw/strafe"],
        ],
        col_widths=[2.2, 1.6, 3.4])
    doc.add_paragraph()

    # ── 3. Architecture Overview ─────────────────────────────────────────────
    add_heading(doc, "3.  Architecture Overview", 1)
    add_para(doc,
        "The architecture is organised into five layers that map cleanly onto the project "
        "execution phases. Each layer is independent enough to be tested in isolation, which "
        "is important on a constrained timeline where debugging one layer at a time is the "
        "only practical strategy.")

    # ── Architecture Diagram ─────────────────────────────────────────────────
    add_heading(doc, "3.1  System Architecture Diagram", 2, "2C3E50")
    add_para(doc, "Figure 1 — Full system architecture showing all five layers and data flows.", italic=True, size=9)

    diag_buf = make_architecture_diagram()
    doc.add_picture(diag_buf, width=Inches(6.8))
    last_para = doc.paragraphs[-1]
    last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # ── 4. Layer Descriptions ─────────────────────────────────────────────────
    add_heading(doc, "4.  Layer Descriptions", 1)

    # Layer 1
    add_heading(doc, "4.1  Layer 1 — Development Infrastructure (Docker)", 2, "2C3E50")
    add_para(doc,
        "All development runs inside a Docker container built on "
        "osrf/ros:melodic-desktop-full (Ubuntu 18.04, ROS 1 Melodic, Gazebo 9). "
        "Host networking (network_mode: host) is used so that ROS topic discovery "
        "and TensorBoard HTTP access work without port-mapping complexity. "
        "An X11 socket and .Xauthority cookie mount enables the Gazebo GUI on the "
        "host display.")
    add_table(doc,
        ["Component", "Role", "Key Detail"],
        [
            ["Docker image",      "Isolated runtime",         "osrf/ros:melodic-desktop-full + deadsnakes Python 3.8"],
            ["docker-compose.yml","Container lifecycle",      "host network · X11 · rw volume → rospug_research/"],
            ["entrypoint.sh",     "ROS environment",          "sources /opt/ros/melodic/setup.bash + catkin devel"],
            ["Python 3.8",        "RL stack runtime",         "stable-baselines3 · gymnasium · tensorboard · rospkg"],
            ["TensorBoard",       "Training monitoring",      "logs/ → port 6006 on host browser"],
            ["Volumes",           "Persistent artefacts",     "logs/ · checkpoints/ written to host filesystem"],
        ],
        col_widths=[1.6, 1.8, 3.8])
    doc.add_paragraph()

    # Layer 2
    add_heading(doc, "4.2  Layer 2 — Simulation (Gazebo 9 / ROS Melodic)", 2, "1565C0")
    add_para(doc,
        "The ROSPug URDF/XACRO model is spawned into an empty Gazebo world via "
        "roslaunch pug_description gazebo.launch. The model exposes 12 position-"
        "controlled joints through ros_control's effort_controllers/JointPositionController, "
        "each with individually tuned PID gains defined in gazebo_control.yaml. "
        "Because the ROSPug URDF does not include a Gazebo IMU plugin, body orientation "
        "(roll and pitch) is obtained from /gazebo/model_states instead of a "
        "sensor_msgs/Imu topic.")

    add_table(doc,
        ["ROS Topic / Service", "Type", "Rate", "Consumer"],
        [
            ["/pug/joint_states",           "sensor_msgs/JointState",   "100 Hz",  "RosPugEnv observation builder"],
            ["/gazebo/model_states",         "gazebo_msgs/ModelStates",  "~100 Hz", "Body roll/pitch (quaternion→RPY)"],
            ["/pug/{joint}_pos.../command",  "std_msgs/Float64",         "50 Hz",   "12 PID controllers (action output)"],
            ["/gazebo/reset_world",          "std_srvs/Empty (service)", "on reset","RosPugEnv.reset()"],
            ["/gazebo/set_physics_properties","gazebo_msgs service",     "on reset","Domain Randomization module"],
        ],
        col_widths=[2.5, 1.8, 0.8, 2.6])
    doc.add_paragraph()

    add_para(doc,
        "The gait baseline controller (sim_gait_controller_v3.py) subscribes to "
        "/cmd_vel and implements a joint-space trot gait with yaw differentials and "
        "hip-abduction strafe. It is not used during RL training but serves as the "
        "non-RL comparison baseline for the real-hardware stretch goal.")

    # Layer 3
    add_heading(doc, "4.3  Layer 3 — RL Environment (RosPugEnv)", 2, "6A1B9A")
    add_para(doc,
        "RosPugEnv (rl_env/rospug_env.py) wraps the ROS/Gazebo stack as a "
        "standard gymnasium.Env so that Stable-Baselines3 can train on it without "
        "knowing anything about ROS. The environment is single-threaded with a fixed "
        "50 Hz control loop. ROS callbacks run in a background thread via rospy's "
        "subscriber mechanism; shared state is protected by a threading.Lock.")

    add_table(doc,
        ["Property", "Value", "Notes"],
        [
            ["Action space",     "Box(−0.5, 0.5, (12,), float32)", "Residual offsets from stand pose (all zeros)"],
            ["Observation space","Box(−∞, ∞, (26,), float32)",     "12 joint pos + 12 joint vel + roll + pitch"],
            ["Control rate",     "50 Hz (DT = 0.02 s)",            "sleep(DT) per step — wall-clock synchronised"],
            ["Max episode steps","500 steps = 10 s",               "Truncation (not termination)"],
            ["Termination cond.","‌|roll| > 0.7 rad OR |pitch| > 0.7 rad", "≈ 40° — falls detected here"],
            ["Reset settle",     "1.0 s physics + 0.5 s stand",    "Avoids carry-over physics state"],
            ["Body orientation", "Inline quat→RPY (no tf dep.)",   "_quat_to_rpy() from /gazebo/model_states"],
        ],
        col_widths=[1.9, 2.6, 2.7])
    doc.add_paragraph()

    add_heading(doc, "Reward Function", 3, "6A1B9A")
    add_para(doc,
        "The reward is designed to encourage forward locomotion while penalising "
        "energy waste and falling. The alive bonus (+0.05/step) breaks the stand-still "
        "local optimum that would otherwise dominate early training:")
    add_para(doc,
        "    r(t) = 3.0 × vx  −  0.001 × Σ(joint_vel²)  +  0.05  −  10.0 × fallen",
        bold=True, size=10)
    add_para(doc,
        "At 0.2 m/s forward velocity: r ≈ +0.65/step. "
        "Standing still: r = +0.05/step. "
        "Falling: r = −10.0 (terminal). "
        "Random actions baseline: mean −1.06/step (dominated by energy penalty).")

    # Layer 4
    add_heading(doc, "4.4  Layer 4 — PPO Training Pipeline", 2, "E65100")
    add_para(doc,
        "Both policies use Proximal Policy Optimization (PPO) from Stable-Baselines3 "
        "with an MLP policy network (two hidden layers of 256 units, tanh activations). "
        "The only difference between Policy A and Policy B is the environment they train "
        "on: Policy A uses a fixed-physics RosPugEnv; Policy B uses a RosPugEnv whose "
        "reset() method re-samples friction, mass, and servo latency before each episode.")

    add_table(doc,
        ["Hyperparameter", "Value", "Rationale"],
        [
            ["Policy network",   "MlpPolicy  [256, 256]",  "Two hidden layers; tanh activation"],
            ["n_steps",          "2048",                    "~4 full 500-step episodes per rollout update"],
            ["batch_size",       "256",                     "Must divide n_steps evenly"],
            ["learning rate",    "3 × 10⁻⁴",               "Adam; reduce to 1×10⁻⁴ if approx_kl > 0.05"],
            ["γ (discount)",     "0.99",                    "Long episode horizon (500 steps)"],
            ["λ (GAE)",          "0.95",                    "Standard GAE-Lambda"],
            ["clip_range",       "0.2",                     "PPO clipping parameter"],
            ["ent_coef",         "0.01",                    "Entropy bonus for exploration"],
            ["Total timesteps",  "500k – 1M",               "~2.6–5.1 hours at ~194k steps/hour"],
            ["Checkpoint every", "10,000 steps",            "Saved to checkpoints/ppo_rospug_*_steps.zip"],
        ],
        col_widths=[1.8, 1.8, 3.6])
    doc.add_paragraph()

    # Domain Randomization
    add_heading(doc, "4.5  Domain Randomization Design (Policy B)", 2, "C62828")
    add_para(doc,
        "Domain Randomization is the key technique that makes Policy B more robust "
        "than Policy A. At the start of each episode, the environment samples new "
        "physics parameters from the distributions below and applies them via Gazebo "
        "services before the robot is reset to its stand pose. The policy must learn "
        "to walk across all of these conditions simultaneously — this forces it to find "
        "a more general locomotion strategy rather than over-fitting to a single set "
        "of physics constants.")

    add_table(doc,
        ["Parameter", "Nominal", "Randomization Range", "Gazebo Mechanism"],
        [
            ["Ground friction (μ)",  "0.8",      "U[0.4, 1.2]",       "/gazebo/set_physics_properties + SDF <mu1>/<mu2>"],
            ["Body mass",            "nominal",  "± 15 % of nominal",  "/gazebo/set_physics_properties body inertial"],
            ["Servo command latency","0 ms",     "U[0, 50] ms",        "step() action buffer with random delay"],
            ["(future) Surface tilt","0°",       "± 5° incline",       "Gazebo ground plane SDF rotation"],
        ],
        col_widths=[1.8, 1.0, 1.8, 3.0])
    doc.add_paragraph()

    add_para(doc,
        "Implementation note: randomization ranges are deliberately conservative. "
        "If Policy B fails to learn at all, the first fix is to narrow the ranges "
        "(e.g. friction U[0.6, 1.0]), not to change the policy architecture. "
        "A documented range-narrowing decision is a legitimate, reportable finding.")

    # Layer 5
    add_heading(doc, "4.6  Layer 5 — Evaluation Framework", 2, "2E7D32")
    add_para(doc,
        "evaluate_policy.py runs both saved policies on held-out physics conditions "
        "that were in neither policy's training distribution (for Policy A: everything "
        "except the nominal values; for Policy B: the tails and extremes outside its "
        "sampled range). This is the project's core scientific result.")

    add_table(doc,
        ["Held-Out Condition", "Friction μ", "Mass offset", "Latency", "Purpose"],
        [
            ["HO-1 (very low friction)",  "0.30",  "nominal",  "0 ms",   "Outside DR lower bound"],
            ["HO-2 (low friction)",       "0.50",  "nominal",  "0 ms",   "Edge of DR range"],
            ["HO-3 (nominal)",            "0.80",  "nominal",  "0 ms",   "Training condition for A"],
            ["HO-4 (high friction)",      "1.10",  "nominal",  "0 ms",   "Edge of DR range"],
            ["HO-5 (heavy + latency)",    "0.80",  "+20 %",    "60 ms",  "Multi-param out-of-dist."],
        ],
        col_widths=[2.1, 1.0, 1.1, 0.8, 2.7])
    doc.add_paragraph()

    add_para(doc,
        "Output metrics per condition: fall rate (%), mean distance walked (m), "
        "mean episode length (steps). Final output: a side-by-side bar chart and "
        "mean ± std summary table comparing Policy A vs. Policy B across all "
        "five held-out conditions.")

    # ── 5. Data Flow ──────────────────────────────────────────────────────────
    add_heading(doc, "5.  Data Flow Summary", 1)
    add_para(doc,
        "The following sequence describes one training step (applies to both policies):")
    bullets = [
        "PPO calls env.reset()  →  /gazebo/reset_world  →  DR module re-samples physics  →  robot placed at stand pose  →  settle 1.0 s",
        "PPO calls env.step(action[12])  →  12 × Float64 published to /pug/{joint}_pos_controller/command",
        "Gazebo steps physics one DT (0.02 s)  →  /pug/joint_states + /gazebo/model_states updated",
        "RosPugEnv reads joint states (pos, vel) + model quaternion  →  assembles obs[26]",
        "Reward computed: r = 3.0 × vx − 0.001 × Σ(vel²) + 0.05 − 10.0 × fallen",
        "If |roll| > 0.7 or |pitch| > 0.7: terminated=True, episode ends",
        "If step_count ≥ 500: truncated=True, episode ends",
        "PPO collects (obs, action, reward, done) tuples into a rollout buffer (n_steps=2048)",
        "After n_steps transitions: PPO computes GAE advantages, runs mini-batch gradient updates",
        "Checkpoint saved every 10,000 steps; TensorBoard scalar logged every update",
    ]
    for b in bullets:
        add_bullet(doc, b)
    doc.add_paragraph()

    # ── 6. File Structure ──────────────────────────────────────────────────────
    add_heading(doc, "6.  Key File Structure", 1)
    add_table(doc,
        ["File / Directory", "Layer", "Purpose"],
        [
            ["Dockerfile",                          "L1",  "Container build: ROS Melodic + Python 3.8 + RL stack"],
            ["docker-compose.yml",                  "L1",  "Container lifecycle, X11, volumes, host networking"],
            ["entrypoint.sh",                       "L1",  "Sources ROS + catkin environments at container start"],
            ["ROSPug/src/pug_description/",         "L2",  "URDF/XACRO model, Gazebo launch, PID controller config"],
            ["ROSPug/src/pug_description/config/\ngazebo_control.yaml", "L2", "12 PID effort controllers definition"],
            ["rospug_research/sim_gait_controller_v3.py", "L2", "Reference trot gait — baseline for real-hardware comparison"],
            ["rl_env/rospug_env.py",                "L3",  "RosPugEnv — gymnasium.Env wrapping Gazebo/ROS1"],
            ["rl_env/__init__.py",                  "L3",  "Package marker; exports RosPugEnv"],
            ["scripts/test_env_random.py",          "L3",  "20-episode random-action sanity test (Step 3 exit criterion)"],
            ["scripts/train_ppo.py",                "L4",  "PPO training script (Policy A; DR variant planned for Policy B)"],
            ["scripts/evaluate_policy.py",          "L5",  "Load checkpoint, run N episodes, print metrics"],
            ["logs/ppo_rospug/",                    "L4",  "TensorBoard event files (Policy A)"],
            ["logs/ppo_rospug_dr/",                 "L4",  "TensorBoard event files (Policy B — planned)"],
            ["checkpoints/",                        "L4",  "PPO .zip checkpoints every 10k steps"],
            ["requirements_rl.txt",                 "L1",  "Pinned RL package versions"],
        ],
        col_widths=[2.6, 0.5, 4.1])
    doc.add_paragraph()

    # ── 7. Interfaces ──────────────────────────────────────────────────────────
    add_heading(doc, "7.  Key Interface Contracts", 1)
    add_para(doc, "These contracts must not change between layers — any change here "
             "requires updating all consumers.", size=9, italic=True)
    add_table(doc,
        ["Interface", "Defined In", "Contract"],
        [
            ["Joint ordering (12)",     "rospug_env.py  JOINT_ORDER",    "rf_j, rf_th, rf_ca, lf_j, lf_th, lf_ca, rb_j, rb_th, rb_ca, lb_j, lb_th, lb_ca"],
            ["Action bounds",           "rospug_env.py  ACTION_LIMIT",   "± 0.5 rad residual from stand pose (all zeros)"],
            ["Observation shape",       "rospug_env.py  obs_space",      "(26,): indices 0-11 pos, 12-23 vel, 24 roll, 25 pitch"],
            ["Control rate",            "rospug_env.py  CTRL_RATE",      "50 Hz (DT=0.02 s) — must match Gazebo physics step"],
            ["Fall threshold",          "rospug_env.py  FALL_THRESH",    "0.7 rad ≈ 40° for both roll and pitch"],
            ["DR parameter service",    "Gazebo ROS API",                "/gazebo/set_physics_properties (before each episode)"],
        ],
        col_widths=[1.8, 1.8, 3.6])
    doc.add_paragraph()

    # ── 8. Stretch Goal ────────────────────────────────────────────────────────
    add_heading(doc, "8.  Stretch Goal — Real Hardware Deployment", 1)
    add_para(doc,
        "If Steps 4–6 are completed ahead of schedule, Policy B is deployed to the "
        "physical ROSPug. The deployment path is:")
    bullets2 = [
        "Export: load Policy B checkpoint with SB3's PPO.load(); run inference (policy.predict(obs)) in a ROS1 Python node",
        "Same action space: publish 12 × Float64 to the exact same /pug/{joint}_position_controller/command topics",
        "Observation: read from /pug/joint_states (pos/vel) and /pug/imu/data or an alternative body-orientation source",
        "Safety protocol: 2–3 s burst tests first; support harness; reduce max joint angle to ± 0.4 rad initially",
        "Comparison: Policy B vs. sim_gait_controller_v3 on a surface not used in simulation (rug, slight incline)",
        "If Policy B fails on hardware: document the reality gap; the simulation result (Step 6) stands independently",
    ]
    for b in bullets2:
        add_bullet(doc, b)
    doc.add_paragraph()

    # ── 9. Limitations & Future Work ───────────────────────────────────────────
    add_heading(doc, "9.  Known Limitations & Future Work", 1)
    add_table(doc,
        ["Limitation", "Impact", "Future Work"],
        [
            ["No IMU Gazebo plugin",       "Body orientation from /gazebo/model_states (unavailable on real robot)", "Add IMU plugin to URDF; or use robot_localization on hardware"],
            ["Single-CPU Gazebo",          "~194k training steps/hour — slower than vectorised simulators",          "Isaac Gym / MJX / parallel Gazebo instances"],
            ["No reality-gap estimator",   "DR ranges are hand-tuned, not auto-calibrated from real trials",         "Bayesian optimisation of DR ranges using real deployment data"],
            ["Fixed MLP policy",           "No temporal context; may struggle with partially-observable states",     "Add LSTM layer or frame-stacking (obs history)"],
            ["No perception / mapping",    "Policies navigate by open-loop locomotion only",                         "Add depth camera or LiDAR; integrate with Nav2/move_base"],
            ["2-week timeline",            "No real-hardware result guaranteed",                                     "Extend to 4-week sprint; add hardware in loop from week 2"],
        ],
        col_widths=[1.9, 2.3, 2.9])
    doc.add_paragraph()

    # ── 10. References ─────────────────────────────────────────────────────────
    add_heading(doc, "10.  References", 1)
    refs = [
        "OpenAI (2019). Solving Rubik's Cube with a Robot Hand (Dactyl). arXiv:1910.07113.",
        "Hwangbo et al. (2019). Learning Agile and Dynamic Motor Skills for Legged Robots (ANYmal). Science Robotics.",
        "Tobin et al. (2017). Domain Randomization for Transferring Deep Neural Networks from Simulation to the Real World. IROS.",
        "Schulman et al. (2017). Proximal Policy Optimization Algorithms. arXiv:1707.06347.",
        "Raffin et al. (2021). Stable-Baselines3: Reliable Reinforcement Learning Implementations. JMLR.",
        "ROSPug Repository: https://github.com/Hiwonder/ROSpug (Jetson_nano_ros1 branch).",
        "ROSPug 2-Week Roadmap: rospug_research/ROSPug_2Week_Roadmap.md",
        "Project Progress Tracker: rospug_research/PROGRESS.md",
    ]
    for r in refs:
        add_bullet(doc, r)

    doc.save(output_path)
    print(f"[OK] Document saved: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == '__main__':
    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "ROSPug_Architecture_Sim2Real_DR.docx")
    build_document(out)
